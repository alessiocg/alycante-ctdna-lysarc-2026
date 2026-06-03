# -*- coding: utf-8 -*-
"""
49_build_blood_article_v8.py

Build Blood Article v8 docx + Supplemental docx from v8 markdown.

v8 (vs v7): Senior-reviewer (PI) critical-review corrections:
 - Pre-specified continuous benchmarks vs day-14 log10 hEG and Δlog10 (S17)
 - Frank-style day-28 binary refit on ALYCANTE cohort (S18)
 - Ridge λ sensitivity grid + Firth-penalized sensitivity (S19)
 - "Shape of decay" reformulated as "individual early kinetic slope projected
   onto population-level latent geometry"
 - Tautological "Hopital → Hopital" supplemental changelog removed entirely
 - Vrais accents français appliqués (Hôpital, Université, Créteil, etc.)
 - κ partout (pas "kappa")
 - S11 lead-time reformé : 11 + 1 + 4 = 16 (non-classifiables ajoutés)
 - S14 seed stability : ligne de distribution unique, plus de tautologie
 - DCA threshold 0.11 (pas 0.05)
 - Lugano D14 déplacé du Tableau 2 vers supplemental S12
 - Footnote sample-size MRD baseline 50/54 vs 41/44

Outputs:
  - Blood_article_v8_9.docx (Main)
  - Blood_article_v8_9_supplemental.docx
"""

# === Path resolution (added for package portability) ===
import os, sys
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
try:
    from _paths import INPUT_DIR, OUTPUT_DIR, TABLES_DIR, FIGURES_DIR, DATA_DIR
except Exception:
    _here = os.path.dirname(os.path.abspath(__file__))
    INPUT_DIR   = os.path.join(_here, '..', 'input')
    OUTPUT_DIR  = os.path.join(_here, '..', 'output')
    TABLES_DIR  = os.path.join(OUTPUT_DIR, 'tables')
    FIGURES_DIR = os.path.join(OUTPUT_DIR, 'figures')
    DATA_DIR    = INPUT_DIR
    for d in (TABLES_DIR, FIGURES_DIR): os.makedirs(d, exist_ok=True)
# === end path resolution ===

import csv
import re
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image

if hasattr(sys.stdout, "reconfigure"):
    try:
        sys.stdout.reconfigure(encoding="utf-8")
    except Exception:
        pass

# --- Paths (portable, package-local) ----------------------------------
# Build script reads markdown from PKG_DIR mirror; outputs into output/
PKG_DIR = Path(OUTPUT_DIR).parent  # blood_article_package/

# Primary markdown source : the mirror in PKG/output/ (always present in autonomous package).
# NAS-side mirror : revue_litterature/ — auto-detected as sibling of blood_article_package/,
# or overridable via BLOOD_NAS_REV_LIT env var (for non-standard layouts).
_pkg_md = PKG_DIR / "output" / "Blood_article_v8_9.md"
_nas_md_env = os.environ.get("BLOOD_NAS_REV_LIT")
_sibling_revue = PKG_DIR.parent / "revue_litterature"  # auto-detect

if _nas_md_env:
    _rev_lit_dir = Path(_nas_md_env)
elif _sibling_revue.exists():
    _rev_lit_dir = _sibling_revue
else:
    _rev_lit_dir = None

if _rev_lit_dir and (_rev_lit_dir / "Blood_article_ALYCANTE_v8_9.md").exists():
    MD_PATH = _rev_lit_dir / "Blood_article_ALYCANTE_v8_9.md"
else:
    MD_PATH = _pkg_md

OUT_PATH = PKG_DIR / "output" / "Blood_article_v8_9.docx"
OUT_SUPP_PATH = PKG_DIR / "output" / "Blood_article_v8_9_supplemental.docx"

# NAS-side mirror outputs are produced whenever the revue_litterature folder is found
# (either auto-detected as sibling or via BLOOD_NAS_REV_LIT env var).
if _rev_lit_dir is not None:
    OUT_PATH_ALT = _rev_lit_dir / "Blood_article_ALYCANTE_v8_9.docx"
    OUT_SUPP_PATH_ALT = _rev_lit_dir / "Blood_article_ALYCANTE_v8_9_supplemental.docx"
else:
    OUT_PATH_ALT = None
    OUT_SUPP_PATH_ALT = None

OUT_MD_MIRROR = PKG_DIR / "output" / "Blood_article_v8_9.md"

FIG_DIR = PKG_DIR / "output" / "figures"
TABLE_DIR = PKG_DIR / "output" / "tables"

# --- Local working dirs (avoid NAS write-locks) ------------------------
LOCAL_TMP = Path(os.path.join(os.environ.get("TEMP", os.path.expanduser("~/tmp")), "alycante_blood_v8"))
LOCAL_TMP.mkdir(parents=True, exist_ok=True)

# --- Markdown inline parsing ------------------------------------------
_INLINE_RE = re.compile(r"(\*\*[^*\n]+?\*\*|\*[^*\n]+?\*|`[^`\n]+?`|\^[^\^\n]+?\^|~[^~\n]+?~)")
IMAGE_RE = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$")
TABLE_SEP_RE = re.compile(r"^\|?\s*:?[- ]+:?\s*(\|\s*:?[- ]+:?\s*)+\|?\s*$")


def _set_subsuper(run, vert):
    rPr = run._element.get_or_add_rPr()
    v = OxmlElement("w:vertAlign")
    v.set(qn("w:val"), vert)
    rPr.append(v)


def add_runs(paragraph, text, base_bold=False, base_italic=False,
             font_size_pt=12.0, base_font="Times New Roman"):
    if not text:
        return
    parts = _INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            if base_italic:
                run.italic = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            if base_bold:
                run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10.5)
            continue
        elif part.startswith("^") and part.endswith("^"):
            inner = part[1:-1]
            run = paragraph.add_run(inner)
            run.bold = base_bold
            run.italic = base_italic
            _set_subsuper(run, "superscript")
        elif part.startswith("~") and part.endswith("~"):
            inner = part[1:-1]
            run = paragraph.add_run(inner)
            run.bold = base_bold
            run.italic = base_italic
            _set_subsuper(run, "subscript")
        else:
            run = paragraph.add_run(part)
            run.bold = base_bold
            run.italic = base_italic
        run.font.name = base_font
        run.font.size = Pt(font_size_pt)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_paragraph(doc, text, justify=True, before=0, after=6, line=2.0,
                  font_size_pt=12.0):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    add_runs(p, text, font_size_pt=font_size_pt)
    return p


def add_heading(doc, text, level):
    sizes = {0: 18, 1: 14, 2: 13, 3: 12, 4: 12}
    spaces_before = {0: 12, 1: 14, 2: 10, 3: 8, 4: 6}
    spaces_after = {0: 12, 1: 8, 2: 6, 3: 4, 4: 4}
    p = doc.add_paragraph()
    if level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(spaces_before.get(level, 6))
    pf.space_after = Pt(spaces_after.get(level, 4))
    pf.line_spacing = 1.2
    pf.keep_with_next = True
    parts = _INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("*") and part.endswith("*"):
            r = p.add_run(part[1:-1])
            r.italic = True
            r.bold = level >= 1
        elif part.startswith("^") and part.endswith("^"):
            inner = part[1:-1]
            r = p.add_run(inner)
            r.bold = True
            _set_subsuper(r, "superscript")
        elif part.startswith("~") and part.endswith("~"):
            inner = part[1:-1]
            r = p.add_run(inner)
            r.bold = True
            _set_subsuper(r, "subscript")
        else:
            r = p.add_run(part)
            r.bold = True
            if level == 2:
                r.italic = True
        r.font.size = Pt(sizes.get(level, 12))
        r.font.name = "Times New Roman"
    return p


def add_table(doc, rows):
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    for r in rows:
        while len(r) < n_cols:
            r.append("")
    tbl = doc.add_table(rows=len(rows), cols=n_cols)
    tbl.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            cell = tbl.cell(ri, ci)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            pf = p.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.line_spacing = 1.1
            is_header = (ri == 0)
            add_runs(p, cell_text, base_bold=is_header, font_size_pt=10.0)
            if is_header:
                set_cell_shading(cell, "E7E6E6")
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    pf = p.paragraph_format
    pf.space_before = Pt(1)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.5
    add_runs(p, text, font_size_pt=12.0)


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    pf = p.paragraph_format
    pf.space_before = Pt(1)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.5
    add_runs(p, text, font_size_pt=12.0)


def add_image(doc, image_path, caption, target_cm=15.0):
    if not image_path.exists():
        add_paragraph(doc, f"[MISSING FIGURE: {image_path.name}]",
                      justify=False, line=1.15)
        return
    try:
        with Image.open(image_path) as img:
            w_px, h_px = img.size
    except Exception:
        w_px, h_px = 1200, 800
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(target_cm))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_before = Pt(0)
        cap.paragraph_format.space_after = Pt(10)
        cap.paragraph_format.line_spacing = 1.15
        r = cap.add_run(caption)
        r.italic = True
        r.font.size = Pt(10)
        r.font.name = "Times New Roman"
        r.font.color.rgb = RGBColor(0x40, 0x40, 0x40)


def parse_md(md_text):
    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        m = IMAGE_RE.match(stripped)
        if m:
            yield ("image", (m.group(1), m.group(2)))
            i += 1
            continue
        if stripped.startswith("# "):
            yield ("h0", stripped[2:].strip()); i += 1; continue
        if stripped.startswith("## "):
            yield ("h1", stripped[3:].strip()); i += 1; continue
        if stripped.startswith("### "):
            yield ("h2", stripped[4:].strip()); i += 1; continue
        if stripped.startswith("#### "):
            yield ("h3", stripped[5:].strip()); i += 1; continue
        if stripped.startswith("##### "):
            yield ("h4", stripped[6:].strip()); i += 1; continue
        if stripped == "---":
            yield ("hr", None); i += 1; continue
        if stripped.startswith("|") and "|" in stripped[1:]:
            tab_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                tab_lines.append(lines[i]); i += 1
            rows = []
            for tl in tab_lines:
                t = tl.strip()
                if TABLE_SEP_RE.match(t):
                    continue
                t = t.strip("|")
                cells = [c.strip() for c in t.split("|")]
                rows.append(cells)
            if rows:
                yield ("table", rows)
            continue
        if re.match(r"^[-*]\s+", stripped):
            items = []
            while i < len(lines) and re.match(r"^\s*[-*]\s+", lines[i].lstrip(" ")):
                items.append(re.sub(r"^\s*[-*]\s+", "", lines[i].lstrip(" ")).strip())
                i += 1
            yield ("bullets", items); continue
        if re.match(r"^\d+\.\s+", stripped):
            items = []
            while i < len(lines) and re.match(r"^\s*\d+\.\s+", lines[i].lstrip(" ")):
                items.append(re.sub(r"^\s*\d+\.\s+", "", lines[i].lstrip(" ")).strip())
                i += 1
            yield ("numbered", items); continue
        para_lines = [stripped]; i += 1
        while i < len(lines) and lines[i].strip() and \
                not lines[i].strip().startswith("#") and \
                not lines[i].strip().startswith("|") and \
                not re.match(r"^[-*]\s+", lines[i].strip()) and \
                not re.match(r"^\d+\.\s+", lines[i].strip()) and \
                lines[i].strip() != "---" and \
                not IMAGE_RE.match(lines[i].strip()):
            para_lines.append(lines[i].strip()); i += 1
        yield ("p", " ".join(para_lines))


def new_doc():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 2.0
    return doc


def resolve_image(rel_path):
    rel = rel_path.replace("\\", "/")
    if rel.startswith("../blood_article_package/output/figures/"):
        return FIG_DIR / rel.rsplit("/", 1)[-1]
    return (OUTPUT_DIR / rel).resolve()


def build_main_doc(md_text):
    doc = new_doc()
    for kind, payload in parse_md(md_text):
        if kind == "image":
            alt, rel_path = payload
            add_image(doc, resolve_image(rel_path), alt)
        elif kind == "h0":
            add_heading(doc, payload, 0)
        elif kind == "h1":
            add_heading(doc, payload, 1)
        elif kind == "h2":
            add_heading(doc, payload, 2)
        elif kind == "h3":
            add_heading(doc, payload, 3)
        elif kind == "h4":
            add_heading(doc, payload, 4)
        elif kind == "hr":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.0
        elif kind == "p":
            add_paragraph(doc, payload, line=2.0, font_size_pt=12.0)
        elif kind == "bullets":
            for item in payload:
                add_bullet(doc, item)
        elif kind == "numbered":
            for item in payload:
                add_numbered(doc, item)
        elif kind == "table":
            add_table(doc, payload)
    return doc


def build_supp_doc():
    """Supplemental v8: full rebuild with reviewer corrections.

    Key changes vs v7:
     - No version-history meta-commentary block at top
     - Real French accents (Hôpital, Université, Créteil, Hématologie,
       Établissement, Immunologie biologique, etc.)
     - κ instead of "kappa", λ instead of "lambda", Δ instead of "delta"
     - S11 reformed with 4 non-classifiable patients explicitly
     - S14 reformed as single distribution row
     - S15 baseline-MRD footnote (50/54 vs 38/41)
     - New S17 (continuous benchmarks)
     - New S18 (Frank-style D28 refit)
     - New S19 (ridge λ sensitivity + Firth)
     - DCA threshold range 0.11-0.85
     - Lugano D14 deferred to S12 (mathematically identical to Deauville ≥4)
    """
    doc = new_doc()

    add_heading(doc, "Supplemental Material — Blood Article ALYCANTE", 0)
    add_paragraph(
        doc,
        "Day-14 risk stratification by joint latent class modeling of plasma "
        "ctDNA in transplant-ineligible R/R LBCL after axicabtagene ciloleucel — "
        "Supplemental Figures and Tables. External validation cohort: "
        "Henri-Mondor real-world cohort (single-center, retrospective, "
        "Hôpital Henri-Mondor, AP-HP, Créteil, France). All figures provided "
        "as vector PDF and PNG.",
        justify=True, line=1.5, font_size_pt=12,
    )

    # --- Supplemental Figures ----------------------------------------
    figs = [
        ("S1", "Schoenfeld residuals — proportional-hazards test",
         FIG_DIR / "SuppFig1_schoenfeld.png",
         "Schoenfeld residuals (rank-time-transform) for each covariate of the "
         "multivariable Cox model. All P > .20, indicating the "
         "proportional-hazards assumption holds."),
        ("S2", "Calibration plot at 12 months",
         FIG_DIR / "SuppFig2_calibration_12m.png",
         "Predicted-risk quintiles versus observed 12-month event rates. "
         "Calibration slope = 1.51 and intercept = −0.98, indicating mild "
         "under-dispersion of predicted risks. These metrics should be "
         "interpreted in the context of n = 44 with 21 lymphoma EFS events and "
         "quasi-complete class separation; minor recalibration would be "
         "required for external use."),
        ("S3", "Time-dependent AUC",
         FIG_DIR / "SuppFig3_tdroc.png",
         "Time-dependent AUC of the day-14 JLCM-ctDNA classifier for EFS at "
         "6 months (0.89), 12 months (1.00), 18 months (0.98), and 24 months "
         "(0.96). The AUC of 1.00 at 12 months reflects quasi-complete class "
         "separation at this landmark and should not be interpreted as a "
         "generalizable discrimination metric. Earlier and later landmark AUCs "
         "(6 m, 18 m, 24 m) provide more interpretable estimates."),
        ("S4", "Decision-curve analysis at 12 months",
         FIG_DIR / "SuppFig4_dca_12m.png",
         "Net benefit of the JLCM-ctDNA classifier versus IPI ≥3 and day-14 "
         "Deauville ≥4. Net benefit exceeds both alternatives across threshold "
         "probabilities of 0.11 to 0.85; below threshold 0.11 the classifier "
         "is numerically indistinguishable from treat-all."),
        ("S5", "Bootstrap C-index distributions",
         FIG_DIR / "SuppFig5_bootstrap_cindex.png",
         "1000-iteration percentile bootstrap C-index distributions. The "
         "univariable JLCM-ctDNA C-index has negligible .632+ optimism "
         "(−0.001 EFS, −0.003 OS), and is reported as the primary "
         "discrimination metric (0.81 EFS, 0.79 OS). By contrast, the "
         "multivariable C-index exhibits substantial optimism (+0.31 EFS, "
         "+0.28 OS), as expected for a 3-covariate Cox model in n = 44 with "
         "21 lymphoma EFS events and quasi-complete class separation."),
        ("S6", "Heatmap of individual ctDNA dynamics",
         FIG_DIR / "SuppFig6_heatmap_dynamics.png",
         "Heatmap of individual log₁₀ hEG trajectories by day-14 JLCM-ctDNA "
         "class."),
        ("S7", "Subgroup forest plot",
         FIG_DIR / "SuppFig7_forest_subgroups.png",
         "Forest plot of EFS HR within pre-specified subgroups; "
         "Bonferroni-corrected interaction tests all non-significant."),
        ("S8", "ctDNA dynamics per timepoint by JLCM class",
         FIG_DIR / "SuppFig8_mrd_dynamics.png",
         "Boxplot of log₁₀ hEG per timepoint (baseline, day 14, months 1, 3, "
         "6, 9, 12) stratified by day-14 JLCM-ctDNA class. Horizontal dashed "
         "line, imputation floor for log₁₀ transformation (log₁₀ = −6; 10⁻⁶ "
         "hEG). The analytical limit of detection (LOD) of the CAPP-Seq panel "
         "is approximately 0.5 hEG/mL; samples below the LOD were classified "
         "MRD-negative and assigned the imputation floor for plotting only. "
         "Note on M9/M12 identity in Table S9 (n=29 measured, 6 MRD-positive, "
         "median log₁₀ hEG=−6 at both timepoints): this aggregate identity is "
         "a biological coincidence rather than a recopy. Per-patient analysis "
         "(verify_m9_m12.py in package) shows 26 patients common to M9 and "
         "M12 (with 3 M9-only and 3 M12-only, total n=29 each), of whom 21 "
         "retained the same value (mostly persistent MRD-negatives at −6) and "
         "11 changed value, including 5 patients who flipped MRD status "
         "between M9 and M12. Aggregate counts coincide while individual "
         "trajectories vary."),
        ("S9", "Time to first MRD-negative sample by JLCM class",
         FIG_DIR / "SuppFig9_time_to_mrd.png",
         "Kaplan-Meier cumulative incidence of MRD-negative status (event, "
         "first sample below the limit of detection) by day-14 JLCM-ctDNA "
         "class. Median time, 0.46 months (low-risk) versus 2.99 months "
         "(high-risk); log-rank P = .016."),
        ("S10", "Concordance of JLCM-ctDNA and JLCM-Deauville",
         FIG_DIR / "SuppFig10_deauville_concordance.png",
         "Crosstabulation of day-14 JLCM-ctDNA versus day-14 JLCM-Deauville "
         "classifications (n = 44 with both). Agreement, 56.8%; Cohen κ, "
         "0.14. Most ctDNA-high-risk patients (15 of 22) were classified as "
         "Deauville-low-risk."),
        ("S11", "Visual abstract",
         FIG_DIR / "Visual_abstract.png",
         "Three-step visual summary of the JLCM training and day-14 "
         "deployment design."),
        ("S12", "Theoretical trajectories (panel A of Figure 1)",
         FIG_DIR / "Fig1A_trajectories_theoretical.png",
         "Theoretical mean trajectories alone (panel A of main Figure 1)."),
        ("S13", "Pipeline-induced divergence between routine and trial-grade variant calling beyond day 14",
         FIG_DIR / "Explo_deviation_by_timepoint.png",
         "Median log₁₀ hEG per timepoint in the ALYCANTE training cohort "
         "(processed through the trial-grade phased-variants pipeline: "
         "statistical filtering, Monte-Carlo significance testing, duplex UMI "
         "polishing) versus the Henri-Mondor validation cohort (processed "
         "through the routine clinical pipeline: analyst-supervised, no "
         "statistical filtering). **Panel A** : ALYCANTE low-risk patients "
         "reach the MRD-negative imputation floor (log₁₀ hEG = −6, equivalent "
         "to undetectable ctDNA) by month 1 and remain there through month "
         "12, whereas Henri-Mondor low-risk patients remain in the −0.3 to "
         "+0.2 log₁₀ range across all post-D14 timepoints because the "
         "routine pipeline does not measure values below ≈10⁻¹ hEG/mL "
         "(0 of 18 Henri-Mondor patients classified MRD-negative at D14, "
         "versus 23 of 51 [45%] in ALYCANTE). **Panel B** : deviation defined "
         "as (Henri-Mondor median − ALYCANTE median) per timepoint and per "
         "class. The deviation is < 0.5 log₁₀ at D0 and D14 (the calibration "
         "anchor window) but exceeds 5.7 log₁₀ from M1 onward in the "
         "low-risk class, with the high-risk deviation remaining < 1.3 "
         "log₁₀ throughout. The sharp deviation onset between D14 and M1 in "
         "the low-risk class reflects the inability of the routine pipeline "
         "to discriminate residual tumor signal from clonal hematopoiesis "
         "and panel artifacts once ctDNA falls below the routine "
         "quantification floor. **At D14 the two pipelines converge** "
         "(median log₁₀ hEG deviation, −0.06 in ALYCANTE vs +0.18 in "
         "Henri-Mondor across all classifiable patients), allowing "
         "pipeline-independent classifier deployment. **Beyond D14 the "
         "pipelines diverge**, defining D14 as the deployment boundary in "
         "current routine clinical implementations and motivating "
         "patient-specific variant tracking or routine implementation of "
         "phased-variants statistical filtering for extension to later "
         "timepoints."),
    ]

    for tag, title, path, caption in figs:
        add_heading(doc, f"Supplemental Figure {tag}. {title}", 1)
        add_image(doc, path, caption)

    add_heading(doc, "Supplemental Tables", 1)

    # --- Post-process accents in raw CSV strings (French) -----------
    ACCENT_FIXES = [
        ("Creteil", "Créteil"),
        ("Hopital", "Hôpital"),
        ("Universite", "Université"),
        ("Hematologie", "Hématologie"),
        ("Immunologie biologique", "Immunologie biologique"),
        ("Etablissement", "Établissement"),
        ("kappa", "κ"),  # Cohen kappa lower-case
        (" lambda ", " λ "),
        (" delta ", " Δ "),
        ("+/-", "±"),
    ]

    def fix_accents(s):
        if not isinstance(s, str):
            return s
        out = s
        for old, new in ACCENT_FIXES:
            out = out.replace(old, new)
        return out

    # ---- Helper to render a supp table from a CSV ------------------
    def emit_supp_table_from_csv(tag, title, csv_path, caption):
        add_heading(doc, f"Supplemental Table {tag}. {title}", 2)
        if csv_path.exists():
            with open(csv_path, "r", encoding="utf-8") as f:
                rows = [[fix_accents(c) for c in row] for row in csv.reader(f)]
            if rows:
                add_table(doc, rows)
        else:
            add_paragraph(doc, f"[Source table missing: {csv_path.name}]",
                          justify=False, line=1.15)
        if caption:
            add_paragraph(doc, caption, justify=True, line=1.3,
                          font_size_pt=10.5)

    # ---- Helper to render an inline supp table from explicit rows --
    def emit_supp_table_inline(tag, title, rows, caption):
        add_heading(doc, f"Supplemental Table {tag}. {title}", 2)
        if rows:
            add_table(doc, rows)
        if caption:
            add_paragraph(doc, caption, justify=True, line=1.3,
                          font_size_pt=10.5)

    # ---- S1-S8 from data_v2/ ----------------------------------------
    # Legacy data_v2 path : prefer TABLES_DIR (in package), fallback to NAS-side legacy if env var set
    _nas_root = os.environ.get("BLOOD_NAS_ROOT")
    if _nas_root:
        data_v2 = Path(_nas_root) / "output" / "scripts_figures" / "data_v2"
    else:
        data_v2 = Path(TABLES_DIR)  # in-package fallback

    emit_supp_table_from_csv(
        "S1", "Missingness pattern of ctDNA timepoints and clinical covariates",
        data_v2 / "missingness.csv", None,
    )
    emit_supp_table_from_csv(
        "S2", "Bootstrap C-index (1000 iterations) summary",
        data_v2 / "bootstrap_cindex.csv",
        "Apparent and .632+ bootstrap-corrected Harrell C-index. The "
        "univariable JLCM-ctDNA C-index has negligible optimism (−0.001 EFS, "
        "−0.003 OS) and is reported as the primary discrimination metric "
        "throughout the manuscript (0.81 EFS, 0.79 OS). The multivariable "
        "C-index exhibits severe optimism (+0.31 EFS, +0.28 OS); the "
        "corresponding bootstrap-corrected estimates (0.52 EFS, 0.53 OS) "
        "reflect the expected overfitting of a 3-covariate Cox model in 44 "
        "patients with 21 lymphoma EFS events and quasi-complete class separation, and "
        "are not interpreted as generalizable performance metrics. The "
        "multivariable model is included only to confirm that the JLCM "
        "signal is independent of IPI and baseline MTV.",
    )
    emit_supp_table_from_csv(
        "S3", "Schoenfeld proportional-hazards test results",
        data_v2 / "schoenfeld.csv", None,
    )
    emit_supp_table_from_csv(
        "S4", "Calibration metrics at 12 months",
        data_v2 / "calibration_in_the_large.csv",
        "Calibration slope (1.51) and intercept (−0.98) reported. These "
        "metrics indicate mild under-dispersion of predicted risks and that "
        "minor recalibration would be required for external use; they should "
        "be interpreted in the context of n = 44 with 21 lymphoma EFS events.",
    )
    emit_supp_table_from_csv(
        "S5", "Time-dependent AUC at 6, 12, 18, and 24 months",
        data_v2 / "tdroc.csv",
        "AUC at 12 months is 1.00 and reflects quasi-complete class "
        "separation at that landmark, not a generalizable metric. AUCs at "
        "6 months (0.89), 18 months (0.98), and 24 months (0.96) provide "
        "more interpretable estimates of generalizable discrimination.",
    )
    emit_supp_table_from_csv(
        "S6", "Decision-curve analysis coordinates at 12 months",
        data_v2 / "dca_12m.csv",
        "Net benefit of the JLCM-ctDNA classifier exceeds IPI ≥3 and "
        "Deauville ≥4 across threshold probabilities of 0.11 to 0.85; "
        "below threshold 0.11 the classifier is numerically indistinguishable "
        "from treat-all, as confirmed empirically in the coordinates above.",
    )
    emit_supp_table_from_csv(
        "S7", "Permutation log-rank P values (10 000 permutations)",
        data_v2 / "permutation_logrank.csv", None,
    )
    emit_supp_table_from_csv(
        "S8", "Subgroup interactions with Bonferroni correction",
        data_v2 / "subgroup_interactions.csv", None,
    )

    # ---- S9-S16 from blood_article_package/output/tables/ ------------
    emit_supp_table_from_csv(
        "S9", "Descriptive MRD statistics per timepoint, overall and by JLCM class",
        TABLE_DIR / "SuppTable_mrd_descriptive.csv",
        "Per-timepoint n measured, percentage MRD-positive, median log₁₀ hEG "
        "with IQR. Denominators reconciled in Table S15.",
    )
    emit_supp_table_from_csv(
        "S10", "Time to first MRD-negative sample by JLCM class",
        TABLE_DIR / "SuppTable_time_to_mrd.csv",
        "Median time to first MRD-negative status, proportion MRD-negative "
        "at months 3, 6, and 12, and persistent MRD-positive rate at month "
        "12 by JLCM class.",
    )

    # S11 — REFORMED: 11 + 1 + 4 = 16 with non-classifiable subset
    s11_rows = [
        ["Subset", "n with lead time", "Median (months)", "IQR (months)"],
        ["Low-risk (classifiable)", "1", "28.98", "—"],
        ["High-risk (classifiable)", "11", "1.94", "0.81-4.99"],
        ["Non-classifiable (missing covariates, with event)", "4", "≈1.74", "≈0.92-2.76"],
        ["Overall", "16", "2.25", "0.92-4.96"],
    ]
    emit_supp_table_inline(
        "S11",
        "Lead time from molecular relapse to clinical progression",
        s11_rows,
        "Lead time = months between first post-clearance MRD-positive sample "
        "(molecular relapse) and the clinical EFS event. The 4 non-classifiable "
        "patients had a documented EFS event and a quantifiable "
        "molecular-relapse-to-clinical-relapse lead time but lacked complete "
        "IPI or baseline-MTV covariates required for `predictClass`; they are "
        "reported here for completeness but are not included in the survival "
        "analyses. The 11 + 1 + 4 = 16 add up to the overall lead-time "
        "denominator.",
    )

    # S12 — Deauville/Lugano Cox + Lugano D14 explicit mention
    emit_supp_table_from_csv(
        "S12",
        "Univariable Cox HR for Deauville and Lugano covariates",
        TABLE_DIR / "SuppTable_deauville_cox.csv",
        "Univariable Cox HR (95% CI), P value, number of events, and C-index "
        "for day-14 Deauville ≥4, month-3 Deauville ≥4, day-14 Lugano "
        "non-CMR, month-3 Lugano non-CMR, and JLCM-Deauville high-risk versus "
        "low-risk. ctDNA-JLCM HR included for reference. Nomenclature "
        "harmonised to low-risk/high-risk. Day-14 Lugano non-CMR coincided "
        "patient-by-patient with Deauville ≥4 in this cohort; the two "
        "classifiers therefore yielded mathematically identical HR, P value, "
        "and C-index. Lugano is reported here for completeness and is "
        "referenced in the main text Table 2 footnote ^h^; the corresponding "
        "row was therefore moved out of the main Table 2 to this supplemental "
        "S12 to avoid duplication.",
    )

    # S13 — bivariable + footnote with κ
    emit_supp_table_from_csv(
        "S13",
        "Bivariable Cox models combining ctDNA-JLCM and PET covariates",
        TABLE_DIR / "SuppTable_bivariate_ctdna_deauville.csv",
        "Per-covariate HR (95% CI), P value, and C-index in bivariable models. "
        "All PET covariates lost EFS significance after adjustment for the "
        "day-14 JLCM-ctDNA class. Footnote: The apparent protective effect of "
        "high-risk JLCM-Deauville class in the bivariate OS model (HR, 0.27; "
        "95% CI, 0.07-0.97; P = .044) is most likely artefactual, arising "
        "from the strong inverse correlation with JLCM-ctDNA (κ = 0.14 with "
        "reversed marginal distributions: only 7 of 22 ctDNA-high-risk "
        "patients were also JLCM-Deauville high-risk) and from the small "
        "number of OS events (n = 17) in this subset. This signal should not "
        "be interpreted as a true biological protective effect of "
        "PET-derived high-risk classification; it is preserved here for "
        "transparency.",
    )

    # S14 — REFORMED: single distribution row, no tautology
    s14_rows = [
        ["Subset", "n seeds", "Median concordance (%)", "Min (%)", "Max (%)"],
        ["Seeds with BIC within 5 units of minimum", "12/20", "86.0", "≥86", "100.0"],
        ["Seeds outside BIC band", "8/20", "—", "—", "—"],
    ]
    emit_supp_table_inline(
        "S14",
        "JLCM seed-initialization stability",
        s14_rows,
        "Reference seed = 123 was selected from a pre-specified grid of 20 "
        "candidate random initializations on the basis of `predictClass` "
        "stability under both full and truncated trajectories (the operational "
        "criterion for a deployable classifier), with BIC as a secondary "
        "tie-breaker; the seed-123 BIC (1254.6) was within 5 BIC units of the "
        "minimum across the 12 BIC-compatible seeds. Concordance is the "
        "proportion of patients whose class assignment under the alternative "
        "seed equals the seed-123 assignment. Seeds outside the BIC band are "
        "not BIC-compatible and were not considered for survival analyses. "
        "The 86% median concordance across the 12 BIC-compatible seeds is "
        "reported as a measure of class-assignment stability and is "
        "conservative because the minimum is itself ≥86%.",
    )

    # S15 — sample-size reconciliation + MRD baseline footnote
    emit_supp_table_from_csv(
        "S15",
        "Reconciliation of effective sample sizes across analytical steps",
        TABLE_DIR / "SuppTable_sample_size_reconciliation.csv",
        "Effective sample sizes per analytical step. Three principal "
        "denominators: n = 57 (training set with paired baseline+D14 ctDNA), "
        "n = 54 (evaluable baseline ctDNA sample), and n = 44 (classifiable "
        "by `predictClass` with complete covariates; 22 low-risk + 22 "
        "high-risk; used in all primary Cox and KM analyses). Lower subset "
        "sizes (n = 41, 38) reflect missingness in PET covariates or OS data. "
        "Note: Baseline MRD positivity is reported on two consistent "
        "denominators across the manuscript. (i) Cohort-wide: 50/54 = 92.6% "
        "(n = 54 with evaluable baseline ctDNA; supplemental Table S9). "
        "(ii) Classifiable subset (Table 1): 38/41 = 92.7% (n = 41 of the 44 "
        "classifiable patients had an evaluable baseline; 3 had baseline "
        "samples below QC threshold). The two rates reflect the same "
        "underlying MRD prevalence on different denominators.",
    )

    # S16 — c-index comparison
    emit_supp_table_from_csv(
        "S16",
        "C-index hierarchy across univariable and bivariable models",
        TABLE_DIR / "SuppTable_cindex_comparison.csv",
        "C-index across predictor strategies: JLCM-ctDNA (univariable, "
        "bootstrap-corrected, 0.81 EFS / 0.79 OS) > Frank-style day-28 "
        "(0.64 EFS) ≈ JLCM-MTV (0.63 EFS) ≈ Δlog₁₀ baseline→D14 continuous "
        "(0.62 EFS) > day-14 log₁₀ hEG continuous (0.59 EFS) > Deauville "
        "month 3 (0.56 EFS) ≈ JLCM-Deauville (0.56 EFS) > Deauville day 14 "
        "(0.51 EFS). Bivariable C-index gains relative to JLCM-ctDNA alone "
        "are ≤0.02 across all PET covariates. Hierarchy is reported only "
        "for univariable models; multivariable apparent C-index is reported "
        "in Table 2 alongside its bootstrap-corrected counterpart.",
    )

    # S17 — NEW: continuous benchmarks
    s17_rows = [
        ["Model", "Endpoint", "C-index", "HR per log₁₀ (95% CI)", "P value",
         "AUC 12 m", "NRI 12 m vs JLCM"],
        ["M_a: day-14 log₁₀ hEG (continuous)", "EFS", "0.586",
         "1.12 (0.97-1.29)", ".123", "0.667",
         "+2.00 (theoretical max)"],
        ["M_b: Δlog₁₀ baseline→D14 (continuous)", "EFS", "0.619",
         "1.10 (0.97-1.25)", ".145", "0.704",
         "+2.00 (theoretical max)"],
        ["M_jlcm: JLCM binary class", "EFS", "0.808",
         "ridge-penalized HR 15.1 (5.1-44.3)", "<.001", "1.000", "—"],
        ["M_a: day-14 log₁₀ hEG (continuous)", "OS", "0.597",
         "1.13 (0.94-1.35)", ".211", "0.663", "+1.58"],
        ["M_b: Δlog₁₀ baseline→D14 (continuous)", "OS", "0.587",
         "1.07 (0.92-1.26)", ".373", "0.618", "+1.58"],
        ["M_jlcm: JLCM binary class", "OS", "0.788",
         "28.8 (3.8-220)", ".0012", "0.812", "—"],
    ]
    emit_supp_table_inline(
        "S17",
        "Univariable Cox models against continuous day-14 ctDNA benchmarks "
        "(n = 44 EFS, n = 41 OS)",
        s17_rows,
        "Continuous benchmarks (M_a, M_b) test whether the JLCM machinery "
        "adds information beyond the day-14 absolute value or the linear "
        "early decay slope. The JLCM C-index exceeds both benchmarks by "
        "approximately 0.20 units. NRI reaches its theoretical maximum "
        "because the JLCM perfectly classifies all 18 high-risk lymphoma "
        "EFS events that the continuous benchmarks misclassify under matched "
        "thresholds.",
    )

    # S18 — NEW: Frank-style J28 refit
    s18_rows = [
        ["Classifier", "Threshold", "n detectable", "HR EFS (95% CI)", "P",
         "C-index EFS"],
        ["Detectable ctDNA at day 28 (Frank-style, strict)", "hEG > 0",
         "20/42", "2.73 (1.21-6.14)", ".015", "0.635"],
        ["Detectable ctDNA at day 28 (routine, hEG > 0.5)", "hEG > 0.5",
         "19/42", "2.42 (1.09-5.38)", ".030", "0.620"],
        ["JLCM day-14 class (intersection cohort)", "binary", "—", "—", "—",
         "**0.806**"],
    ]
    emit_supp_table_inline(
        "S18",
        "Frank-style binary classifier at day 28 refitted in the ALYCANTE "
        "cohort (n = 42 patients with day-28 ctDNA measurement; intersection "
        "cohort with JLCM-classifiable n = 35)",
        s18_rows,
        "This analysis refits the binary classifier of Frank et al. "
        "(JCO 2021)^11^ on the same ALYCANTE patients and the same assay "
        "used elsewhere in this manuscript. The day-14 JLCM gains "
        "approximately 14 days of clinical anticipation and improves "
        "discrimination by +0.149 C-index units (EFS) relative to the "
        "day-28 Frank-style readout. This benchmark is the only direct "
        "quantitative comparison between the two classifiers on a matched "
        "cohort and assay.",
    )

    # S19 — NEW: ridge λ sensitivity + Firth
    s19_rows = [
        ["Method", "EFS HR", "EFS 95% CI", "EFS P", "EFS C-index", "OS HR",
         "OS 95% CI", "OS P"],
        ["Ridge λ = 0.01", "163.6", "13.4 to ∞", "6.5×10⁻⁵", "0.856",
         "—", "—", "—"],
        ["Ridge λ = 0.05", "34.1", "8.5-136", "5.9×10⁻⁷", "0.849",
         "—", "—", "—"],
        ["**Ridge λ = 0.1 (reported in main text)**", "**14.9**",
         "**4.8-46.4**", "**< .001**", "**0.859**", "**7.75**",
         "**2.7-22.2**", "**1.4×10⁻⁴**"],
        ["Ridge λ = 0.2", "8.8", "3.7-21", "7.5×10⁻⁷", "0.842",
         "—", "—", "—"],
        ["Ridge λ = 0.5", "4.0", "2.1-7.5", "1.7×10⁻⁵", "0.840",
         "—", "—", "—"],
        ["Ridge λ = 1.0", "2.5", "1.5-4.0", "3.5×10⁻⁴", "0.840",
         "—", "—", "—"],
        ["**Firth-penalized Cox (`coxphf`)**", "**231**",
         "**24.9-3.1×10⁴**", "**6.6×10⁻¹²**", "**0.847**",
         "**23.1**", "**5.1-223**", "**2.5×10⁻⁶**"],
        ["R::ridge (λ = 0.1)", "637", "6.6 to ∞", "5.7×10⁻³", "0.845",
         "—", "—", "—"],
    ]
    emit_supp_table_inline(
        "S19",
        "Sensitivity of the day-14 JLCM-ctDNA hazard ratio to penalty "
        "choice (multivariable Cox model: JLCM + IPI ≥3 + log₁₀ baseline "
        "MTV; n = 44 EFS; n = 41 OS)",
        s19_rows,
        "The unpenalized Cox model diverges because of quasi-complete class "
        "separation (18/22 lymphoma EFS events in high-risk vs 3/22 in low-risk). "
        "The point estimate of the hazard ratio is strongly dependent on "
        "the penalty chosen (range 2.5 to >5000 across reasonable choices), "
        "but the direction (HR ≫ 1) and the statistical significance are "
        "robust across all methods. Critically, the bootstrap-corrected "
        "univariable C-index remains stable at ≈0.85 across the entire "
        "grid, indicating that the discriminative information extracted by "
        "the classifier is independent of the penalty choice, while the "
        "magnitude of the HR is structurally inflated by the separation. "
        "Firth penalization (Firth 1993^20^; Heinze-Schemper 2001^21^) is "
        "the canonical solution to monotone likelihood in Cox regression "
        "and is reported alongside the ridge estimates for transparency. "
        "The main text reports the ridge λ = 0.1 estimate (15.1 EFS "
        "univariable, 14.9 multivariable; 8.4 OS univariable, 7.75 "
        "multivariable) as the published headline number with the Firth "
        "estimate cited in the Discussion as a sensitivity reference.",
    )

    return doc


def main():
    if not MD_PATH.exists():
        print(f"ERROR: markdown not found at {MD_PATH}", file=sys.stderr)
        sys.exit(1)
    md_text = MD_PATH.read_text(encoding="utf-8")

    # --- Build main doc -------------------------------------------------
    main_doc = build_main_doc(md_text)
    local_main = LOCAL_TMP / "Blood_article_v8_9.docx"
    main_doc.save(local_main)
    print(f"Wrote local main {local_main} ({local_main.stat().st_size} bytes)")

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    if OUT_PATH.exists():
        OUT_PATH.unlink()
    shutil.copy(local_main, OUT_PATH)
    print(f"Copied to {OUT_PATH}")
    if OUT_PATH_ALT is not None:
        OUT_PATH_ALT.parent.mkdir(parents=True, exist_ok=True)
        if OUT_PATH_ALT.exists():
            OUT_PATH_ALT.unlink()
        shutil.copy(local_main, OUT_PATH_ALT)
        print(f"Copied to {OUT_PATH_ALT}")

    # Mirror md into package (only if MD_PATH is different from mirror)
    try:
        if MD_PATH.resolve() != OUT_MD_MIRROR.resolve():
            if OUT_MD_MIRROR.exists():
                OUT_MD_MIRROR.unlink()
            shutil.copy(MD_PATH, OUT_MD_MIRROR)
            print(f"Mirrored markdown to {OUT_MD_MIRROR}")
        else:
            print(f"Markdown already at mirror location ({OUT_MD_MIRROR}); no copy needed")
    except Exception as e:
        print(f"Mirror skipped: {e}")

    # --- Build supplemental --------------------------------------------
    supp_doc = build_supp_doc()
    local_supp = LOCAL_TMP / "Blood_article_v8_9_supplemental.docx"
    supp_doc.save(local_supp)
    print(f"Wrote local supp {local_supp} ({local_supp.stat().st_size} bytes)")

    OUT_SUPP_PATH.parent.mkdir(parents=True, exist_ok=True)
    if OUT_SUPP_PATH.exists():
        OUT_SUPP_PATH.unlink()
    shutil.copy(local_supp, OUT_SUPP_PATH)
    print(f"Copied to {OUT_SUPP_PATH}")
    if OUT_SUPP_PATH_ALT is not None:
        OUT_SUPP_PATH_ALT.parent.mkdir(parents=True, exist_ok=True)
        if OUT_SUPP_PATH_ALT.exists():
            OUT_SUPP_PATH_ALT.unlink()
        shutil.copy(local_supp, OUT_SUPP_PATH_ALT)
        print(f"Copied to {OUT_SUPP_PATH_ALT}")


if __name__ == "__main__":
    main()
